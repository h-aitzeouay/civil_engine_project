from __future__ import annotations

from pathlib import Path
from typing import Any


def fmt(value: Any, ndigits: int = 2) -> str:
    if value is None:
        return "-"

    try:
        if isinstance(value, float):
            return f"{value:.{ndigits}f}"
        return str(value)
    except Exception:
        return "-"


def yes_no(value: bool) -> str:
    return "OUI" if value else "NON"


def export_project_summary_docx(
    dashboard: dict[str, Any],
    output_path: str | Path,
) -> str:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    output_path = Path(output_path)

    doc = Document()

    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(9)

    title = doc.add_heading("RAPPORT DE SYNTHESE PROJET FONDATIONS", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Synthèse de livraison — INGENIERIE.COM STRUCTURAL AI")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    project_name = dashboard.get("project_name", "-")
    ready = bool(dashboard.get("ready_to_deliver", False))
    status = dashboard.get("status", "-")

    doc.add_heading("1. Statut global", level=1)
    doc.add_paragraph(f"Projet : {project_name}")
    doc.add_paragraph(f"Statut global : {status}")
    doc.add_paragraph(f"Prêt à livrer : {yes_no(ready)}")

    blocking = dashboard.get("blocking_reasons", [])

    if blocking:
        doc.add_paragraph("Points bloquants :")
        for item in blocking:
            doc.add_paragraph(f"- {item}")
    else:
        doc.add_paragraph("Aucun point bloquant détecté par le contrôle automatique.")

    doc.add_heading("2. Statuts techniques", level=1)

    global_status = dashboard.get("global_status", {})

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Contrôle"
    table.rows[0].cells[1].text = "Statut"

    for key in ["quality", "reinforcement", "punching", "anchorage"]:
        row = table.add_row().cells
        row[0].text = key
        row[1].text = str(global_status.get(key, "-"))

    doc.add_heading("3. Fondations finales", level=1)

    foundations = dashboard.get("summaries", {}).get("foundations", {})

    doc.add_paragraph(f"Nombre de fondations finales : {foundations.get('count', 0)}")
    doc.add_paragraph(f"Répartition par type : {foundations.get('by_type', {})}")
    doc.add_paragraph(f"Groupes d'épaisseurs H : {foundations.get('thickness_groups', {})}")

    doc.add_heading("4. Corrections appliquées", level=1)

    remediation = dashboard.get("remediation_summary", {})

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Type de correction"
    table.rows[0].cells[1].text = "Nombre"

    rows = [
        ("Corrections géométriques / emprise", remediation.get("geometry_corrections_count", 0)),
        ("Corrections poinçonnement / épaisseur H", remediation.get("punching_corrections_count", 0)),
        ("Corrections ancrages / crosses 135°", remediation.get("anchorage_corrections_count", 0)),
    ]

    for key, value in rows:
        row = table.add_row().cells
        row[0].text = key
        row[1].text = str(value)

    doc.add_paragraph("")
    doc.add_paragraph("Détail des corrections principales :")

    for item in remediation.get("punching_corrections", []):
        fid = item.get("foundation_id", "-")
        old_h = item.get("H_old_m", "-")
        new_h = item.get("H_new_m", "-")
        util = item.get("utilization_before", "-")
        doc.add_paragraph(f"- {fid} : H {old_h} m → {new_h} m, taux initial {util}")

    doc.add_heading("5. Métré principal", level=1)

    boq = dashboard.get("summaries", {}).get("boq", {})

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Poste"
    table.rows[0].cells[1].text = "Quantité"

    boq_rows = [
        ("Béton fondations", f"{fmt(boq.get('concrete_m3'))} m³"),
        ("Béton de propreté", f"{fmt(boq.get('clean_concrete_m3'))} m³"),
        ("Coffrage latéral", f"{fmt(boq.get('formwork_m2'))} m²"),
        ("Acier fondations", f"{fmt(boq.get('foundation_steel_kg'))} kg"),
        ("Acier attentes", f"{fmt(boq.get('starter_steel_kg'))} kg"),
        ("Acier total", f"{fmt(boq.get('total_steel_kg'))} kg"),
    ]

    for key, value in boq_rows:
        row = table.add_row().cells
        row[0].text = key
        row[1].text = value

    doc.add_heading("6. Livrables", level=1)

    deliverables = dashboard.get("deliverables", {})

    for key, endpoint in deliverables.items():
        doc.add_paragraph(f"- {key} : {endpoint}")

    doc.add_heading("7. Notes d'ingénierie", level=1)

    for note in dashboard.get("engineering_notes", []):
        doc.add_paragraph(f"- {note}")

    doc.add_paragraph("")
    doc.add_paragraph("Réserves :")
    doc.add_paragraph(
        "- Ce rapport est une synthèse automatique. "
        "La validation finale par un ingénieur structure reste obligatoire avant exécution."
    )
    doc.add_paragraph(
        "- Les longueurs d'ancrage, recouvrements, rayons de cintrage, dispositions sismiques "
        "et efforts définitifs doivent être confirmés dans la note finale."
    )

    doc.save(output_path)
    return str(output_path)


def export_project_summary_pdf(
    dashboard: dict[str, Any],
    output_path: str | Path,
) -> str:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

    output_path = Path(output_path)

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=1.4 * cm,
        leftMargin=1.4 * cm,
        topMargin=1.4 * cm,
        bottomMargin=1.4 * cm,
    )

    styles = getSampleStyleSheet()
    story = []

    def h1(text: str) -> None:
        story.append(Paragraph(text, styles["Heading1"]))
        story.append(Spacer(1, 0.15 * cm))

    def h2(text: str) -> None:
        story.append(Paragraph(text, styles["Heading2"]))
        story.append(Spacer(1, 0.10 * cm))

    def p(text: str) -> None:
        story.append(Paragraph(str(text), styles["Normal"]))
        story.append(Spacer(1, 0.08 * cm))

    def table_2cols(rows: list[list[Any]]) -> None:
        data = [["Poste", "Valeur"]] + [[str(a), str(b)] for a, b in rows]
        table = Table(data, colWidths=[7.0 * cm, 9.0 * cm])
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(table)
        story.append(Spacer(1, 0.25 * cm))

    h1("RAPPORT DE SYNTHESE PROJET FONDATIONS")
    p("Synthèse de livraison — INGENIERIE.COM STRUCTURAL AI")

    project_name = dashboard.get("project_name", "-")
    ready = bool(dashboard.get("ready_to_deliver", False))
    status = dashboard.get("status", "-")

    h2("1. Statut global")
    table_2cols([
        ["Projet", project_name],
        ["Statut global", status],
        ["Prêt à livrer", yes_no(ready)],
    ])

    blocking = dashboard.get("blocking_reasons", [])

    if blocking:
        p("Points bloquants :")
        for item in blocking:
            p(f"- {item}")
    else:
        p("Aucun point bloquant détecté par le contrôle automatique.")

    h2("2. Statuts techniques")

    global_status = dashboard.get("global_status", {})

    table_2cols([
        ["Qualité", global_status.get("quality", "-")],
        ["Ferraillage", global_status.get("reinforcement", "-")],
        ["Poinçonnement", global_status.get("punching", "-")],
        ["Ancrages", global_status.get("anchorage", "-")],
    ])

    h2("3. Fondations finales")

    foundations = dashboard.get("summaries", {}).get("foundations", {})

    table_2cols([
        ["Nombre de fondations finales", foundations.get("count", 0)],
        ["Répartition par type", foundations.get("by_type", {})],
        ["Groupes d'épaisseurs H", foundations.get("thickness_groups", {})],
    ])

    h2("4. Corrections appliquées")

    remediation = dashboard.get("remediation_summary", {})

    table_2cols([
        ["Corrections géométriques", remediation.get("geometry_corrections_count", 0)],
        ["Corrections poinçonnement", remediation.get("punching_corrections_count", 0)],
        ["Corrections ancrages 135°", remediation.get("anchorage_corrections_count", 0)],
    ])

    for item in remediation.get("punching_corrections", []):
        fid = item.get("foundation_id", "-")
        old_h = item.get("H_old_m", "-")
        new_h = item.get("H_new_m", "-")
        util = item.get("utilization_before", "-")
        p(f"- {fid} : H {old_h} m → {new_h} m, taux initial {util}")

    h2("5. Métré principal")

    boq = dashboard.get("summaries", {}).get("boq", {})

    table_2cols([
        ["Béton fondations", f"{fmt(boq.get('concrete_m3'))} m³"],
        ["Béton de propreté", f"{fmt(boq.get('clean_concrete_m3'))} m³"],
        ["Coffrage latéral", f"{fmt(boq.get('formwork_m2'))} m²"],
        ["Acier fondations", f"{fmt(boq.get('foundation_steel_kg'))} kg"],
        ["Acier attentes", f"{fmt(boq.get('starter_steel_kg'))} kg"],
        ["Acier total", f"{fmt(boq.get('total_steel_kg'))} kg"],
    ])

    h2("6. Livrables")

    for key, endpoint in dashboard.get("deliverables", {}).items():
        p(f"- {key} : {endpoint}")

    h2("7. Notes d'ingénierie")

    for note in dashboard.get("engineering_notes", []):
        p(f"- {note}")

    p("Réserve : ce rapport est une synthèse automatique. Validation finale par ingénieur structure obligatoire avant exécution.")
    p("Les longueurs d'ancrage, recouvrements, rayons de cintrage, dispositions sismiques et efforts définitifs doivent être confirmés dans la note finale.")

    doc.build(story)
    return str(output_path)
