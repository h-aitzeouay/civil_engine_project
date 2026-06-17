from __future__ import annotations

from pathlib import Path
from typing import Any


def fmt(value: Any, ndigits: int = 2) -> str:
    if value is None:
        return "-"

    try:
        if isinstance(value, float):
            return f"{value:.{ndigits}f}"
        return str(value)
    except Exception:
        return "-"


def fmt_counts(counts: dict[str, Any]) -> str:
    """Formate {'SE': 10, 'SI': 2} -> "SE : 10, SI : 2" ; vide -> "-"."""
    if not counts:
        return "-"
    return ", ".join(f"{key} : {value}" for key, value in counts.items())


def export_calculation_report_docx(
    report: dict[str, Any],
    output_path: str | Path,
) -> str:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    output_path = Path(output_path)

    doc = Document()

    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(9)

    title = doc.add_heading("NOTE DE CALCUL FONDATIONS", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Prédimensionnement automatique — INGENIERIE.COM")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    project = report.get("project", {})
    hypotheses = report.get("hypotheses", {})

    doc.add_heading("1. Identification", level=1)
    doc.add_paragraph(f"Projet : {project.get('name', '-')}")
    doc.add_paragraph(f"Phase : {project.get('phase', '-')}")
    doc.add_paragraph(f"Méthode : {report.get('method', '-')}")

    doc.add_heading("2. Hypothèses", level=1)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Paramètre"
    table.rows[0].cells[1].text = "Valeur"

    rows = [
        ("Contrainte admissible sol", f"{fmt(hypotheses.get('q_allowable_kPa'))} kPa"),
        ("fck", f"{fmt(hypotheses.get('fck_mpa'))} MPa"),
        ("fyk", f"{fmt(hypotheses.get('fyk_mpa'))} MPa"),
        ("gamma_s", fmt(hypotheses.get("gamma_s"))),
        ("gamma_c", fmt(hypotheses.get("gamma_c"))),
        ("Enrobage", f"{fmt(hypotheses.get('cover_m'))} m"),
        ("Béton de propreté", f"{fmt(hypotheses.get('clean_concrete_m'))} m"),
        ("Diamètre principal", f"HA{fmt(hypotheses.get('phi_main_mm'), 0)}"),
        ("Attentes poteaux", f"HA{fmt(hypotheses.get('starter_diameter_mm'), 0)}"),
    ]

    for key, value in rows:
        row = table.add_row().cells
        row[0].text = key
        row[1].text = value

    doc.add_heading("3. Stratégie de fondations", level=1)

    strategy_summary = report.get("strategy_summary", {})
    doc.add_paragraph(f"Nombre de fondations finales : {strategy_summary.get('foundation_count', 0)}")
    doc.add_paragraph(f"Répartition par type : {fmt_counts(strategy_summary.get('by_type', {}))}")
    doc.add_paragraph(f"Avertissements : {strategy_summary.get('warnings_count', 0)}")

    table = doc.add_table(rows=1, cols=8)
    table.style = "Table Grid"

    headers = ["ID", "Type", "Poteaux", "A", "B", "H", "qELS", "Statut"]

    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    for f in report.get("foundation_strategy", {}).get("final_foundations", []):
        row = table.add_row().cells
        row[0].text = str(f.get("id", "-"))
        row[1].text = str(f.get("type", "-"))
        row[2].text = ",".join(f.get("columns", []))
        row[3].text = fmt(f.get("A_m"))
        row[4].text = fmt(f.get("B_m"))
        row[5].text = fmt(f.get("H_m"))
        row[6].text = fmt(f.get("soil_pressure_ELS_kPa"))
        row[7].text = str(f.get("status", "-"))

    # --- Descente de charges par semelle ---
    lt = report.get("load_takedown", {})
    doc.add_heading("4. Descente de charges par semelle", level=1)
    doc.add_paragraph(f"Combinaison ELU appliquée : {lt.get('elu_combination', '1.35G + 1.5Q')}")
    doc.add_paragraph("Combinaison ELS : G + Q. N_ELS = Gk + Qk (cumul de tous les niveaux portés).")
    doc.add_paragraph("Les surcharges locales (calque CHARGE-Q) sont incluses dans Qk le cas échéant.")

    lt_table = doc.add_table(rows=1, cols=8)
    lt_table.style = "Table Grid"
    lt_headers = ["Semelle", "Poteau", "Niveaux", "SGk (kN)", "SQk (kN)", "N_ELS (kN)", "N_ELU (kN)", "Comb. ELU"]
    for i, header in enumerate(lt_headers):
        lt_table.rows[0].cells[i].text = header
    for r in lt.get("by_footing", []):
        row = lt_table.add_row().cells
        row[0].text = str(r.get("foundation_id", "-"))
        row[1].text = str(r.get("column_id", "-"))
        row[2].text = ",".join(r.get("levels_supported", [])) or "-"
        row[3].text = fmt(r.get("sum_Gk_kN"))
        row[4].text = fmt(r.get("sum_Qk_kN"))
        row[5].text = fmt(r.get("N_ELS_kN"))
        row[6].text = fmt(r.get("N_ELU_kN"))
        row[7].text = str(r.get("combination_elu", "-"))

    tot = lt.get("totals", {})
    if tot:
        doc.add_paragraph(
            f"Totaux projet — SGk : {fmt(tot.get('total_Gk_kN'))} kN | "
            f"SQk : {fmt(tot.get('total_Qk_kN'))} kN | "
            f"N_ELS : {fmt(tot.get('total_N_ELS_kN'))} kN | "
            f"N_ELU : {fmt(tot.get('total_N_ELU_kN'))} kN"
        )

    doc.add_heading("5. Vérification du ferraillage", level=1)

    reinf_summary = report.get("reinforcement_summary", {})
    doc.add_paragraph(f"Statut global : {reinf_summary.get('status', '-')}")
    doc.add_paragraph(f"Fondations vérifiées : {reinf_summary.get('foundations_checked', 0)}")
    doc.add_paragraph(f"Répartition : {fmt_counts(reinf_summary.get('by_status', {}))}")
    doc.add_paragraph(f"Avertissements : {reinf_summary.get('warnings_count', 0)}")
    doc.add_paragraph(f"Erreurs : {reinf_summary.get('errors_count', 0)}")

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    headers = ["Fondation", "Type", "rho_l réel", "Statut"]

    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    for c in report.get("reinforcement", {}).get("final_check", {}).get("checks", []):
        row = table.add_row().cells
        row[0].text = str(c.get("foundation_id", "-"))
        row[1].text = str(c.get("foundation_type", "-"))
        row[2].text = fmt(c.get("rho_l_real_for_punching"), 5)
        row[3].text = str(c.get("status", "-"))

    doc.add_heading("6. Poinçonnement final", level=1)

    punching_summary = report.get("punching_summary", {})
    doc.add_paragraph(f"Statut global : {punching_summary.get('status', '-')}")
    doc.add_paragraph(f"Utilisation maximale : {fmt(punching_summary.get('worst_utilization'), 3)}")
    doc.add_paragraph(f"Fondation critique : {punching_summary.get('worst_foundation', '-')}")

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    headers = ["Fondation", "Type", "rho utilisé", "Utilisation", "Statut"]

    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    for c in report.get("punching", {}).get("checks", []):
        row = table.add_row().cells
        row[0].text = str(c.get("foundation_id", "-"))
        row[1].text = str(c.get("foundation_type", "-"))
        row[2].text = fmt(c.get("rho_l_real_used"), 5)
        row[3].text = fmt(c.get("worst_utilization"), 3)
        row[4].text = str(c.get("status", "-"))

    doc.add_heading("7. Ancrages et recouvrements", level=1)

    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"

    headers = ["Poteau", "Fondation", "Attentes", "Lbd", "L0", "Retour", "Forme"]

    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    for item in report.get("anchorage", {}).get("rows", []):
        starter = item.get("starter_bars", {})
        anchorage = item.get("anchorage", {})

        row = table.add_row().cells
        row[0].text = str(item.get("column_id", "-"))
        row[1].text = str(item.get("foundation_id", "-"))
        row[2].text = str(starter.get("label", "-"))
        row[3].text = fmt(anchorage.get("Lbd_m"))
        row[4].text = fmt(anchorage.get("lap_L0_m"))
        row[5].text = fmt(anchorage.get("hook_leg_m"))
        row[6].text = str(anchorage.get("recommended_shape", "-"))

    doc.add_heading("8. Métré estimatif", level=1)

    totals = report.get("boq_summary", {})

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Poste"
    table.rows[0].cells[1].text = "Quantité"

    _beams_c = float(totals.get("beams_concrete_m3", 0.0) or 0.0)
    _tot_c = float(totals.get("concrete_m3", 0.0) or 0.0)
    boq_rows = [
        ("Béton semelles", f"{fmt(round(_tot_c - _beams_c, 2))} m³"),
        ("Béton poutres (chaînage/PR/liaison)", f"{fmt(_beams_c)} m³"),
        ("Béton total", f"{fmt(_tot_c)} m³"),
        ("Béton de propreté", f"{fmt(totals.get('clean_concrete_m3'))} m³"),
        ("Coffrage latéral semelles", f"{fmt(totals.get('formwork_m2'))} m²"),
        ("Acier semelles", f"{fmt(totals.get('foundation_steel_kg'))} kg"),
        ("Acier attentes", f"{fmt(totals.get('starter_steel_kg'))} kg"),
        ("Acier poutres", f"{fmt(totals.get('beams_steel_kg'))} kg"),
        ("Acier total", f"{fmt(totals.get('total_steel_kg'))} kg"),
    ]

    for key, value in boq_rows:
        row = table.add_row().cells
        row[0].text = key
        row[1].text = value

    doc.add_heading("9. Réserves et validations obligatoires", level=1)

    for item in report.get("limitations", []):
        doc.add_paragraph(f"- {item}")

    doc.add_paragraph("")
    doc.add_paragraph("Document généré automatiquement par INGENIERIE.COM STRUCTURAL AI.")
    doc.add_paragraph("Contrôle et validation par ingénieur structure obligatoire avant exécution.")

    doc.save(output_path)

    return str(output_path)


def export_calculation_report_pdf(
    report: dict[str, Any],
    output_path: str | Path,
) -> str:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak

    output_path = Path(output_path)

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=1.3 * cm,
        leftMargin=1.3 * cm,
        topMargin=1.3 * cm,
        bottomMargin=1.3 * cm,
    )

    styles = getSampleStyleSheet()
    story = []

    def h1(text: str) -> None:
        story.append(Paragraph(text, styles["Heading1"]))
        story.append(Spacer(1, 0.15 * cm))

    def h2(text: str) -> None:
        story.append(Paragraph(text, styles["Heading2"]))
        story.append(Spacer(1, 0.10 * cm))

    def p(text: str) -> None:
        story.append(Paragraph(str(text), styles["Normal"]))
        story.append(Spacer(1, 0.08 * cm))

    def small_table(headers: list[str], rows: list[list[Any]]) -> None:
        data = [headers] + [[str(v) for v in row] for row in rows]

        table = Table(data, repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 7),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))

        story.append(table)
        story.append(Spacer(1, 0.25 * cm))

    project = report.get("project", {})
    hypotheses = report.get("hypotheses", {})

    h1("NOTE DE CALCUL FONDATIONS")
    p("Prédimensionnement automatique — INGENIERIE.COM")
    p(f"Projet : {project.get('name', '-')}")
    p(f"Phase : {project.get('phase', '-')}")
    p(f"Méthode : {report.get('method', '-')}")

    h2("1. Hypothèses")

    small_table(
        ["Paramètre", "Valeur"],
        [
            ["Contrainte admissible sol", f"{fmt(hypotheses.get('q_allowable_kPa'))} kPa"],
            ["fck", f"{fmt(hypotheses.get('fck_mpa'))} MPa"],
            ["fyk", f"{fmt(hypotheses.get('fyk_mpa'))} MPa"],
            ["gamma_s", fmt(hypotheses.get("gamma_s"))],
            ["gamma_c", fmt(hypotheses.get("gamma_c"))],
            ["Enrobage", f"{fmt(hypotheses.get('cover_m'))} m"],
            ["Béton de propreté", f"{fmt(hypotheses.get('clean_concrete_m'))} m"],
            ["Diamètre principal", f"HA{fmt(hypotheses.get('phi_main_mm'), 0)}"],
            ["Attentes poteaux", f"HA{fmt(hypotheses.get('starter_diameter_mm'), 0)}"],
        ],
    )

    h2("2. Stratégie de fondations")

    strategy_summary = report.get("strategy_summary", {})
    p(f"Nombre de fondations finales : {strategy_summary.get('foundation_count', 0)}")
    p(f"Répartition par type : {fmt_counts(strategy_summary.get('by_type', {}))}")
    p(f"Avertissements : {strategy_summary.get('warnings_count', 0)}")

    foundation_rows = []

    for f in report.get("foundation_strategy", {}).get("final_foundations", []):
        foundation_rows.append([
            f.get("id"),
            f.get("type"),
            ",".join(f.get("columns", [])),
            fmt(f.get("A_m")),
            fmt(f.get("B_m")),
            fmt(f.get("H_m")),
            fmt(f.get("soil_pressure_ELS_kPa")),
            f.get("status", "-"),
        ])

    small_table(
        ["ID", "Type", "Poteaux", "A", "B", "H", "qELS", "Statut"],
        foundation_rows,
    )

    # --- Descente de charges par semelle ---
    lt = report.get("load_takedown", {})
    h2("3. Descente de charges par semelle")
    p(f"Combinaison ELU appliquee : {lt.get('elu_combination', '1.35G + 1.5Q')}")
    p("Combinaison ELS : G + Q. N_ELS = Gk + Qk (cumul des niveaux portes).")
    p("Surcharges locales (calque CHARGE-Q) incluses dans Qk le cas echeant.")
    lt_rows = []
    for r in lt.get("by_footing", []):
        lt_rows.append([
            r.get("foundation_id", "-"),
            r.get("column_id", "-"),
            ",".join(r.get("levels_supported", [])) or "-",
            fmt(r.get("sum_Gk_kN")),
            fmt(r.get("sum_Qk_kN")),
            fmt(r.get("N_ELS_kN")),
            fmt(r.get("N_ELU_kN")),
            r.get("combination_elu", "-"),
        ])
    small_table(
        ["Semelle", "Poteau", "Niveaux", "SGk (kN)", "SQk (kN)", "N_ELS (kN)", "N_ELU (kN)", "Comb. ELU"],
        lt_rows,
    )
    tot = lt.get("totals", {})
    if tot:
        p(f"Totaux projet — SGk: {fmt(tot.get('total_Gk_kN'))} kN | "
          f"SQk: {fmt(tot.get('total_Qk_kN'))} kN | "
          f"N_ELS: {fmt(tot.get('total_N_ELS_kN'))} kN | "
          f"N_ELU: {fmt(tot.get('total_N_ELU_kN'))} kN")

    h2("4. Ferraillage")

    reinf_summary = report.get("reinforcement_summary", {})
    p(f"Statut global : {reinf_summary.get('status', '-')}")
    p(f"Fondations vérifiées : {reinf_summary.get('foundations_checked', 0)}")
    p(f"Répartition : {fmt_counts(reinf_summary.get('by_status', {}))}")

    reinf_rows = []

    for c in report.get("reinforcement", {}).get("final_check", {}).get("checks", []):
        reinf_rows.append([
            c.get("foundation_id"),
            c.get("foundation_type"),
            fmt(c.get("rho_l_real_for_punching"), 5),
            c.get("status"),
        ])

    small_table(
        ["Fondation", "Type", "rho_l réel", "Statut"],
        reinf_rows,
    )

    h2("5. Poinçonnement final")

    punching_summary = report.get("punching_summary", {})
    p(f"Statut global : {punching_summary.get('status', '-')}")
    p(f"Utilisation maximale : {fmt(punching_summary.get('worst_utilization'), 3)}")
    p(f"Fondation critique : {punching_summary.get('worst_foundation', '-')}")

    punching_rows = []

    for c in report.get("punching", {}).get("checks", []):
        punching_rows.append([
            c.get("foundation_id"),
            c.get("foundation_type"),
            fmt(c.get("rho_l_real_used"), 5),
            fmt(c.get("worst_utilization"), 3),
            c.get("status"),
        ])

    small_table(
        ["Fondation", "Type", "rho", "Utilisation", "Statut"],
        punching_rows,
    )

    story.append(PageBreak())

    h2("6. Ancrages et recouvrements")

    anchorage_rows = []

    for item in report.get("anchorage", {}).get("rows", []):
        starter = item.get("starter_bars", {})
        anchorage = item.get("anchorage", {})

        anchorage_rows.append([
            item.get("column_id"),
            item.get("foundation_id"),
            starter.get("label"),
            fmt(anchorage.get("Lbd_m")),
            fmt(anchorage.get("lap_L0_m")),
            fmt(anchorage.get("hook_leg_m")),
            anchorage.get("recommended_shape"),
        ])

    small_table(
        ["Poteau", "Fondation", "Attentes", "Lbd", "L0", "Retour", "Forme"],
        anchorage_rows,
    )

    h2("7. Métré estimatif")

    totals = report.get("boq_summary", {})

    _beams_c2 = float(totals.get("beams_concrete_m3", 0.0) or 0.0)
    _tot_c2 = float(totals.get("concrete_m3", 0.0) or 0.0)
    small_table(
        ["Poste", "Quantité"],
        [
            ["Béton semelles", f"{fmt(round(_tot_c2 - _beams_c2, 2))} m³"],
            ["Béton poutres (chaînage/PR/liaison)", f"{fmt(_beams_c2)} m³"],
            ["Béton total", f"{fmt(_tot_c2)} m³"],
            ["Béton de propreté", f"{fmt(totals.get('clean_concrete_m3'))} m³"],
            ["Coffrage latéral semelles", f"{fmt(totals.get('formwork_m2'))} m²"],
            ["Acier semelles", f"{fmt(totals.get('foundation_steel_kg'))} kg"],
            ["Acier attentes", f"{fmt(totals.get('starter_steel_kg'))} kg"],
            ["Acier poutres", f"{fmt(totals.get('beams_steel_kg'))} kg"],
            ["Acier total", f"{fmt(totals.get('total_steel_kg'))} kg"],
        ],
    )

    h2("8. Réserves et validations obligatoires")

    for item in report.get("limitations", []):
        p(f"- {item}")

    p("Document généré automatiquement par INGENIERIE.COM STRUCTURAL AI.")
    p("Contrôle et validation par ingénieur structure obligatoire avant exécution.")

    doc.build(story)

    return str(output_path)
